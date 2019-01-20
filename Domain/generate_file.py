from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def generate(missing_links, old_links, place):
    """

    :param missing_links: список источников, на которые пропущены ссылки
    :param old_links: список источников, которые устарели
    :param place:
    :return:
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        section.bottom_margin = Cm(2)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)

    if len(missing_links):
        doc.add_paragraph('На следующие источники отсутствуют ссылки', style=style)
        for line in missing_links:
            paragraph = doc.add_paragraph(line, style=style)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    if len(old_links):
        doc.add_paragraph('Устаревшие источники', style=style)
        for line in old_links:
            paragraph = doc.add_paragraph(line, style=style)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.save(place)
