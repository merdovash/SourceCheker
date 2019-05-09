import re
import os.path
from collections import defaultdict
from datetime import datetime
from typing import List

# Настройки
import docx
from docx.oxml import CT_P, CT_Tbl, CT_Inline
from docx.shape import InlineShape
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.document import Document


MAX_ERRORS = 3  # Макс. кол-во ошибок для остановки поиска  [3]
MIN_LENGTH = 14  # Миним. длина источника   [14]
MIN_STAGE = 2  # Кол-во заголовков списка [2, но иногда может быть 1]


class NoSourcesException(Exception):
    pass


def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph.
    """
    print(type(parent))
    if isinstance(parent, docx.document.Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iter():
        if isinstance(child, docx.oxml.CT_P):
            yield docx.text.paragraph.Paragraph(child, parent)
        elif isinstance(child, docx.oxml.table.CT_Tbl):
            yield Table(child, parent)
        elif isinstance(child, docx.oxml.CT_Inline):
            yield InlineShape(child)


def get_year(source):
    """
    Возвращает год выхода источника
    :param source: str
    :return: int or None
    """

    years = [int(year) for year in re.findall('[1-2][0-9]{3}', source)]
    current_year = datetime.now().year
    years = [x for x in years if x <= current_year]
    if len(years) == 0:
        return None
    if len(years) == 1:
        return years[0]
    return max(years)


def find_missing_src(file_path, callback=lambda x, y: None, min_year=1000):
    """
    TODO возвращает объект с полями:
        - sources: List - список источников
            - text: str - текст источника
            - year: int - год источника
            - has_link: bool - есть ли ссылка на этот источник
            - is_modern: bool - проходит ли проверку на современность
            - paragraphs: List[str] - список параграфов, имеющих ссылки на этот источник
        - author: List[str, str, str] - автор работы
        - year: int - год написания работы

    TODO извлекать даты  с помощью '[1-2][0-9]{3}' => 1000-2999

    :param min_year: int минимальный год
    :param file_path: str путь к файлу
    :param callback: Callable[str, int[0:100]] принимает строку о текущей задаче и число с текущим процентом выполнения
    :return: Tuple[List[str], List[int], List[int]]
        возвращает список всех источников
        и список индексов источников на которые есть ссылки
        и список индексов устаревших источников
    """

    if not os.path.isfile(file_path):
        return None, None
    else:
        from docx import Document
        document: Document = Document(file_path)

    try:
        source_checker = SourceReferenceChecker(document, min_year=min_year)
    except NoSourcesException:
        source_checker = None

    checkers: List[Checker] = [source_checker]

    for paragraph in iter_block_items(document):
        for checker in checkers:
            if checker is not None:
                checker.process(paragraph)

    return checkers


class Checker:
    def process(self, obj: Paragraph or Table or InlineShape, **kwargs):
        raise NotImplementedError()

    def report(self):
        raise NotImplementedError()


class SourceReferenceChecker(Checker):

    def __init__(self, document: Document, **kwargs):
        def find_sources(document):
            """
            Возвращает индекс параграфа в документе, с которого начниатеся список литературы
            :param document: WordDocument
            :return: int
            """
            last_mention = None
            for index, paragraph in enumerate(document.paragraphs):
                paragraph_text = paragraph.text.lower()
                if "список" in paragraph_text:
                    if ("литератур" in paragraph_text) or ("источник" in paragraph_text):
                        last_mention = index

            return last_mention

        def is_source(paragraph):
            """
            Проверяет является ли параграф источником
            :param paragraph: Paragraph
            :return: bool
            """
            text = paragraph.text.lower()
            if len(re.findall('[0-9]+(\s)*[CСcс]\.', text)) > 0:
                # 52 c.
                return True

            if len(re.findall('[CСcс]\.(\s)*[0-9]+', text)):
                # c. 52
                return True

            if len(re.findall('с\.(\s){0,2}[0-9]{1,4}(\s){0,2}-(\s){0,2}[0-9]{1,4}', text)):
                # c. 52 – 57
                return True

            if '//' in paragraph.text:
                return True

            if get_year(paragraph.text):
                return True

            return False

        self.sources = []
        self.sources_re = {}
        self.old = []

        self.links = defaultdict(list)

        paragraphs = document.paragraphs

        sources_header_index = find_sources(document)
        if sources_header_index is None:
            raise NoSourcesException()
        sources_paragraphs = paragraphs[sources_header_index:]

        for index, paragraph in enumerate(sources_paragraphs, sources_header_index):
            if is_source(paragraph):
                self.sources.append(paragraph.text)
                self.sources_re[paragraph] = re.compile(
                    '\[([0-9],\s{0,2}){0,4}' + str(len(self.sources)) + '(,\s{0,2}[0-9]){0,4}\]'
                )
                year = get_year(paragraph.text)
                if year is not None and year < kwargs.get('min_year', 2000):
                    self.old.append(len(self.sources) - 1)

    def process(self, obj: Paragraph or Table, **kwargs):
        if isinstance(obj, Paragraph):
            paragraph_text = obj.text.lower()
            for source, regex in self.sources_re.items():
                res = regex.findall(paragraph_text)
                if len(res):
                    self.links[source.text].append(obj.text)

    def report(self):
        for index, source in enumerate(self.sources, 1):
            print(index, source, ':', self.links[source])


def run():
    while True:

        print("\n")
        input_file_path = input("Path? ")

        args = find_missing_src(input_file_path)
        print(args)

        input_user_action = input("\nOpen another file? (y/n): ")

        if input_user_action.lower() == "n":
            break


if __name__ == "__main__":
    run()
